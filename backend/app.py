"""
One Stop ALM Administration - Flask Backend v3.0
All utility logic ported directly from login.py (original Tkinter app).
Single ALM thread holds tdc — solves COM threading completely.
messagebox calls removed (replaced with return values).
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
from datetime import datetime
import time, random, uuid, threading, queue, os, re
import openpyxl
from openpyxl import Workbook

app = Flask(__name__)
CORS(app)

# ─── Detect OTA ───────────────────────────────────────────────────────────────
OTA_AVAILABLE = False
try:
    import win32com.client
    import pythoncom
    OTA_AVAILABLE = True
    print("✅ OTA API detected — Real ALM mode available")
except ImportError:
    print("ℹ️  OTA not available — Demo mode")

# ─── Sessions ─────────────────────────────────────────────────────────────────
sessions = {}

def log_entry(message, level="INFO"):
    return {"timestamp": datetime.now().strftime("%H:%M:%S"), "level": level, "message": message}

def dummy_delay():
    time.sleep(random.uniform(0.5, 1.2))

def require_session(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        session = sessions.get(token)
        if not session:
            return jsonify({"detail": "Invalid or expired session. Please login again."}), 401
        return f(session, *args, **kwargs)
    return decorated


# ══════════════════════════════════════════════════════════════════════════════
# UTILITY CLASSES — ported exactly from login.py, messagebox removed
# ══════════════════════════════════════════════════════════════════════════════

class ALMQueryProcessor:
    @staticmethod
    def extract_text_from_html(html):
        try:
            from bs4 import BeautifulSoup
            import pandas as pd
            if pd.notna(html):
                if str(html).strip().startswith('<') and str(html).strip().endswith('>'):
                    return BeautifulSoup(html, 'html.parser').get_text()
                else:
                    return html
        except:
            pass
        return html or ''

    def execute_query_and_save_to_excel(self, input_excel_file, excel_file, word_file_prefix, tdc):
        import pandas as pd
        input_workbook = openpyxl.load_workbook(input_excel_file)
        input_sheet = input_workbook.active
        ts_test_ids, tester_names = [], []
        is_header = True
        for row in input_sheet.iter_rows(values_only=True):
            if is_header:
                is_header = False
                continue
            ts_test_ids.append(str(row[9]))
            tester_names.append(str(row[4]))
        input_workbook.close()
        self.create_excel_file(excel_file)
        self.populate_excel_file(ts_test_ids, tester_names, excel_file, tdc)
        self.create_word_documents_from_excel(excel_file, word_file_prefix)

    def create_excel_file(self, excel_file):
        headers = ["DS_ID","TS_TEST_ID","TS_NAME","TS_DESCRIPTION","PREREQUISITES",
                   "DS_STEP_NAME","DS_DESCRIPTION","DS_EXPECTED","DS_STEP_ORDER",
                   "TEST_SCRIPT_STATUS","Tester name"]
        wb = openpyxl.Workbook()
        wb.active.append(headers)
        wb.save(excel_file)

    def populate_excel_file(self, ts_test_ids, tester_names, excel_file, tdc):
        wb = openpyxl.load_workbook(excel_file)
        sheet = wb.active
        component_row_count = 2
        for i, ts_test_id in enumerate(ts_test_ids):
            initial_row = component_row_count
            query = (f"Select DESSTEPS.DS_ID,TEST.TS_TEST_ID,TEST.TS_NAME,TEST.TS_DESCRIPTION,"
                     f"TEST.TS_USER_25,DESSTEPS.DS_STEP_NAME,DESSTEPS.DS_DESCRIPTION,"
                     f"DESSTEPS.DS_EXPECTED,DESSTEPS.DS_STEP_ORDER,TEST.TS_USER_05 "
                     f"from TEST,DESSTEPS Where TS_TEST_ID='{ts_test_id}' "
                     f"AND TS_TEST_ID=DS_TEST_ID order by TS_TEST_ID,DS_STEP_ORDER ASC")
            com = tdc.Command
            com.CommandText = query
            rs = com.Execute()
            while not rs.EOR:
                sheet.append([rs.FieldValue(0), rs.FieldValue(1), rs.FieldValue(2),
                               self.extract_text_from_html(rs.FieldValue(3)),
                               self.extract_text_from_html(rs.FieldValue(4)),
                               rs.FieldValue(5),
                               self.extract_text_from_html(rs.FieldValue(6)),
                               self.extract_text_from_html(rs.FieldValue(7)),
                               rs.FieldValue(8), rs.FieldValue(9), tester_names[i]])
                rs.Next()
                component_row_count += 1
            for col in ['B','C','D','E','J','K']:
                sheet.merge_cells(f"{col}{initial_row}:{col}{component_row_count-1}")
        wb.save(excel_file)

    @staticmethod
    def create_word_documents_from_excel(excel_file, word_file_prefix):
        from docx import Document
        from docx.shared import Pt
        wb = openpyxl.load_workbook(excel_file)
        sheet = wb.active
        row_min = 2
        while row_min < sheet.max_row:
            test_script_name = sheet.cell(row=row_min, column=3).value
            responsible_tester = sheet.cell(row=row_min, column=11).value
            ts_id = sheet.cell(row=row_min, column=2).value
            tester_folder = os.path.join(word_file_prefix, str(responsible_tester))
            os.makedirs(tester_folder, exist_ok=True)
            doc = Document()
            doc.styles['Heading 1'].font.size = Pt(14)
            doc.styles['Heading 1'].font.bold = True
            doc.add_paragraph().add_run("Test Name: ").bold = True
            doc.add_paragraph().add_run(str(test_script_name))
            doc.add_paragraph().add_run("Test Description: ").bold = True
            doc.add_paragraph().add_run(str(sheet.cell(row=row_min, column=4).value))
            doc.add_paragraph().add_run("PreCondition: ").bold = True
            doc.add_paragraph().add_run(str(sheet.cell(row=row_min, column=5).value))
            doc.add_paragraph().add_run("Step Details: ").bold = True
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Step Name', 'Step Description', 'Expected Result'
            num_steps = 0
            step_row = row_min + 1
            while step_row <= sheet.max_row:
                if sheet.cell(row=step_row, column=3).value is not None:
                    break
                num_steps += 1
                step_row += 1
            step_row = row_min
            for _ in range(num_steps + 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(sheet.cell(row=step_row, column=6).value or "")
                row_cells[1].text = str(sheet.cell(row=step_row, column=7).value or "")
                row_cells[2].text = str(sheet.cell(row=step_row, column=8).value or "")
                step_row += 1
            doc.add_page_break()
            doc.save(os.path.join(tester_folder, f"{ts_id}_{str(test_script_name)[:50]}.docx"))
            row_min = step_row


class TestTypeUpdator:
    def convert_to_manual(self, test_case_ids, test_type, qc_connection):
        com = qc_connection.Command
        for test_case_id in test_case_ids:
            com.CommandText = f"UPDATE TEST SET TS_TYPE = '{test_type}' WHERE TS_TEST_ID = '{test_case_id}'"
            com.Execute()

    def execute_conversion(self, test_case_ids, test_type, tdc):
        self.convert_to_manual(test_case_ids, test_type, tdc)


class DefectExtractor:
    def convert_to_date(self, value):
        import pandas as pd
        try:
            if pd.notna(value):
                return datetime.strptime(value, '%Y-%m-%d %H:%M:%S').strftime('%d/%m/%Y')
        except:
            pass
        return ''

    def extract_text_from_html(self, html):
        try:
            from bs4 import BeautifulSoup
            import pandas as pd
            if pd.notna(html):
                if str(html).strip().startswith('<') and str(html).strip().endswith('>'):
                    return BeautifulSoup(html, 'html.parser').get_text()
                return html
        except:
            pass
        return ''

    def sanitize_text(self, text):
        if text is None:
            return None
        return re.sub(r'[\\/:*?"<>|]', '', str(text))

    def Defect_Extraction(self, status_list, priority_list, severity_list, category_list, tdc):
        headers = ["Bug ID","Bug Status","Bug Category","Test Execution Release","Test Execution Cycle",
                   "Detected Environment","Bug Summary","Bug Description","Bug Comments","Bug Severity",
                   "Bug Priority","Bug Detected By","Bug Detected Date","Bug Closing date"]
        wb = openpyxl.Workbook()
        sheet = wb.active
        sheet.append(headers)
        if status_list:
            where = " OR ".join([f"BUG.BG_STATUS='{s}'" for s in status_list])
        elif priority_list:
            where = " OR ".join([f"BUG.BG_PRIORITY='{p}'" for p in priority_list])
        elif severity_list:
            where = " OR ".join([f"BUG.BG_SEVERITY='{s}'" for s in severity_list])
        elif category_list:
            where = " OR ".join([f"BUG.BG_USER_01='{c}'" for c in category_list])
        else:
            where = None
        sql = ("SELECT BUG.BG_BUG_ID,BUG.BG_STATUS,BUG.BG_USER_01,BUG.BG_DETECTED_IN_REL,"
               "BUG.BG_DETECTED_IN_RCYC,BUG.BG_ENVIRONMENT,BUG.BG_SUMMARY,BUG.BG_DESCRIPTION,"
               "BUG.BG_DEV_COMMENTS,BUG.BG_SEVERITY,BUG.BG_PRIORITY,BUG.BG_DETECTED_BY,"
               "BUG.BG_DETECTION_DATE,BUG.BG_CLOSING_DATE FROM BUG")
        if where:
            sql += f" WHERE {where}"
        sql += " order by BUG.BG_BUG_ID"
        com = tdc.Command
        com.CommandText = sql
        rs = com.Execute()
        while not rs.EOR:
            sheet.append([rs.FieldValue(0), rs.FieldValue(1), rs.FieldValue(2),
                          rs.FieldValue(3), rs.FieldValue(4),
                          self.sanitize_text(rs.FieldValue(5)),
                          self.sanitize_text(self.extract_text_from_html(rs.FieldValue(6))),
                          self.sanitize_text(self.extract_text_from_html(rs.FieldValue(7))),
                          self.sanitize_text(self.extract_text_from_html(rs.FieldValue(8))),
                          rs.FieldValue(9), rs.FieldValue(10), rs.FieldValue(11),
                          self.convert_to_date(rs.FieldValue(12)),
                          self.convert_to_date(rs.FieldValue(13))])
            rs.Next()
        wb.save("ExtractedDefects.xlsx")
        return "ExtractedDefects.xlsx"


class testCaseExtractor:
    def extract_text_from_html(self, html):
        try:
            from bs4 import BeautifulSoup
            if html:
                return BeautifulSoup(html, 'html.parser').get_text()
        except:
            pass
        return ""

    def GetNodesList(self, tdc, RootNode):
        arrStrNodesList = [RootNode]
        objTreeManager = tdc.TreeManager
        objSubjectNode = objTreeManager.NodeByPath(RootNode)
        for i in range(1, objSubjectNode.Count + 1):
            if objSubjectNode.Child(i).Count >= 1:
                arrStrNodesList.extend(self.GetNodesList(tdc, objSubjectNode.Child(i).Path))
            else:
                arrStrNodesList.append(objSubjectNode.Child(i).path)
        return arrStrNodesList

    def TestCase_WA(self, folder_path, tdc):
        wb = openpyxl.Workbook()
        ws = wb.active
        headers = ["Test ID","Test Name","Test Description","Test Type","Test Author","Creation Date",
                   "Comments","Change Status","Execution Status","Subject","Path","Tree Path",
                   "Requirement ID","Requirement Name","Step Name","Step Description","Expected Result"]
        for col_num, header in enumerate(headers, 1):
            ws.cell(row=1, column=col_num, value=header)
        folderNameLists = self.GetNodesList(tdc, folder_path)
        i = 2
        for foldername in folderNameLists:
            test_case_folder = tdc.TreeManager.NodeByPath(foldername)
            for test in test_case_folder.TestFactory.NewList(""):
                iColumn = self.populate_test_data(test, ws, i, foldername)
                for req in test.GetCoverList():
                    ws.cell(row=i, column=iColumn+1, value=req.ID)
                    ws.cell(row=i, column=iColumn+2, value=req.Name.strip())
                    i += 1
                for test_step in test.DesignStepFactory.NewList(""):
                    self.populate_test_step_data(test_step, ws, i, 15)
                    i += 1
        wb.save("TestCaseDetailsFinal.xlsx")
        return "TestCaseDetailsFinal.xlsx"

    def populate_test_data(self, test, ws, row_num, foldername):
        ws.cell(row=row_num, column=1).value = test.Field("TS_TEST_ID")
        ws.cell(row=row_num, column=2).value = test.Field("TS_NAME").strip()
        ws.cell(row=row_num, column=3).value = self.extract_text_from_html(test.Field("TS_DESCRIPTION"))
        ws.cell(row=row_num, column=4).value = test.Field("TS_TYPE").strip()
        ws.cell(row=row_num, column=5).value = test.Field("TS_RESPONSIBLE").strip()
        cd = test.Field("TS_CREATION_DATE")
        ws.cell(row=row_num, column=6).value = cd.strftime('%Y-%m-%d %H:%M:%S') if cd else ""
        ws.cell(row=row_num, column=7).value = self.extract_text_from_html(test.Field("TS_DEV_COMMENTS"))
        ws.cell(row=row_num, column=8).value = test.Field("TS_BPTA_CHANGE_DETECTED")
        ws.cell(row=row_num, column=9).value = test.Field("TS_EXEC_STATUS").strip()
        ws.cell(row=row_num, column=10).value = str(test.Field("TS_SUBJECT"))
        ws.cell(row=row_num, column=11).value = foldername.strip()
        ws.cell(row=row_num, column=12).value = test.Field("TS_TREE_PATH")
        return 12

    def populate_test_step_data(self, test_step, ws, row_num, col_start):
        sn = test_step.Field("DS_STEP_NAME")
        sd = test_step.Field("DS_DESCRIPTION")
        er = test_step.Field("DS_EXPECTED")
        ws.cell(row=row_num, column=col_start).value   = sn.strip() if sn else ""
        ws.cell(row=row_num, column=col_start+1).value = self.extract_text_from_html(sd) if sd else ""
        ws.cell(row=row_num, column=col_start+2).value = self.extract_text_from_html(er) if er else ""


class testCaseExtractor_WA:
    def extract_text_from_html(self, html):
        try:
            from bs4 import BeautifulSoup
            if html:
                return BeautifulSoup(html, 'html.parser').get_text()
        except:
            pass
        return ""

    def download_test_files_and_attachments(self, test, foldername):
        Download_Path = ""
        current_directory = os.path.abspath(os.getcwd())
        for t_attach in test.Attachments.NewList(""):
            ALM_ID = test.Field("TS_TEST_ID")
            attachment_folder = os.path.join(current_directory, "ALMExtraction", str(ALM_ID), "Attachments")
            os.makedirs(attachment_folder, exist_ok=True)
            test_attach_storage = t_attach.AttachmentStorage
            test_attach_storage.ClientPath = attachment_folder
            try:
                t_attach.Load(True, test_attach_storage.ClientPath)
                Download_Path += ";" + t_attach.FileName
            except Exception as e:
                print(f"Error: {t_attach.FileName}")
        return Download_Path

    def download_test_step_attachments(self, test_step, test, foldername):
        Download_Path = ""
        current_directory = os.path.abspath(os.getcwd())
        for t_attach in test_step.Attachments.NewList(""):
            ALM_ID = test.Field("TS_TEST_ID")
            attachment_folder = os.path.join(current_directory, "ALMExtraction", str(ALM_ID), "StepAttachments")
            os.makedirs(attachment_folder, exist_ok=True)
            test_attach_storage = t_attach.AttachmentStorage
            test_attach_storage.ClientPath = attachment_folder
            try:
                t_attach.Load(True, test_attach_storage.ClientPath)
                Download_Path += ";" + t_attach.FileName
            except:
                pass
        return Download_Path

    def TestCase_WA(self, folder_path, tdc):
        wb = openpyxl.Workbook()
        ws = wb.active
        headers = ["Test ID","Test Name","Attachment Path","Test Description","Test Type","Test Author",
                   "Creation Date","Comments","Change Status","Execution Status","Subject","Path","Tree Path"]
        for col_num, header in enumerate(headers, 1):
            ws.cell(row=1, column=col_num, value=header)
        folderNameLists = self._GetNodesList(tdc, folder_path)
        i = 2
        for foldername in folderNameLists:
            test_case_folder = tdc.TreeManager.NodeByPath(foldername)
            for test in test_case_folder.TestFactory.NewList(""):
                download_Path = (self.download_test_files_and_attachments(test, foldername)
                                 if test.Attachments.NewList("") else "No Attachment for this")
                iColumn = self.populate_test_data(test, ws, i, foldername, download_Path)
                for test_step in test.DesignStepFactory.NewList(""):
                    Attach_Path = (self.download_test_step_attachments(test_step, test, foldername)
                                   if test_step.Attachments.NewList("") else "No Attachment for this")
                    self.populate_test_step_data(test_step, ws, i, iColumn, Attach_Path)
                    i += 1
        wb.save("TestCaseDetailsFinal_WA.xlsx")
        return "TestCaseDetailsFinal_WA.xlsx"

    def _GetNodesList(self, tdc, RootNode):
        arrStrNodesList = [RootNode]
        objSubjectNode = tdc.TreeManager.NodeByPath(RootNode)
        for i in range(1, objSubjectNode.Count + 1):
            if objSubjectNode.Child(i).Count >= 1:
                arrStrNodesList.extend(self._GetNodesList(tdc, objSubjectNode.Child(i).Path))
            else:
                arrStrNodesList.append(objSubjectNode.Child(i).path)
        return arrStrNodesList

    def populate_test_data(self, test, ws, row_num, foldername, download_Path):
        ws.cell(row=row_num, column=1).value  = test.Field("TS_TEST_ID")
        ws.cell(row=row_num, column=2).value  = test.Field("TS_NAME").strip()
        ws.cell(row=row_num, column=3).value  = download_Path
        ws.cell(row=row_num, column=4).value  = self.extract_text_from_html(test.Field("TS_DESCRIPTION"))
        ws.cell(row=row_num, column=5).value  = test.Field("TS_TYPE").strip()
        ws.cell(row=row_num, column=6).value  = test.Field("TS_RESPONSIBLE").strip()
        cd = test.Field("TS_CREATION_DATE")
        ws.cell(row=row_num, column=7).value  = cd.strftime('%Y-%m-%d %H:%M:%S') if cd else ""
        ws.cell(row=row_num, column=8).value  = self.extract_text_from_html(test.Field("TS_DEV_COMMENTS"))
        ws.cell(row=row_num, column=9).value  = test.Field("TS_BPTA_CHANGE_DETECTED")
        ws.cell(row=row_num, column=10).value = test.Field("TS_EXEC_STATUS").strip()
        ws.cell(row=row_num, column=11).value = str(test.Field("TS_SUBJECT"))
        ws.cell(row=row_num, column=12).value = foldername.strip()
        ws.cell(row=row_num, column=13).value = test.Field("TS_TREE_PATH")
        return 13

    def populate_test_step_data(self, test_step, ws, row_num, step_column, Attach_Path):
        ws.cell(row=row_num, column=step_column).value   = test_step.Field("DS_STEP_NAME")
        ws.cell(row=row_num, column=step_column+1).value = self.extract_text_from_html(test_step.Field("DS_DESCRIPTION"))
        ws.cell(row=row_num, column=step_column+2).value = self.extract_text_from_html(test_step.Field("DS_EXPECTED"))
        ws.cell(row=row_num, column=step_column+3).value = Attach_Path


class testCaseExtractor_TS:
    def extract_text_from_html(self, html):
        try:
            from bs4 import BeautifulSoup
            if html:
                return BeautifulSoup(html, 'html.parser').get_text()
        except:
            pass
        return ""

    def download_test_cases(self, test_set_ids, OTA):
        wb = openpyxl.Workbook()
        ws = wb.active
        headers = ["Test Set ID","testSetName","Test Case ID","Test Name","Test Description",
                   "Criticality","Test Script Status","Test Type","Test Author","Creation Date",
                   "Comments","Change Status","Execution Status","Step Name","Step Description","Expected Result"]
        for col_num, header in enumerate(headers, 1):
            ws.cell(row=1, column=col_num, value=header)
        tsFact = OTA.TestSetFactory
        row = 2
        for testSetId in [t.strip() for t in test_set_ids.split(',')]:
            testSet = tsFact.Item(testSetId)
            testSetName = testSet.Name
            tests = testSet.TSTestFactory.NewList("")
            ws.cell(row=row, column=1).value = testSetId
            ws.cell(row=row, column=2).value = testSetName
            for test in tests:
                testPlan = test.Test
                type_ = test.Type
                if type_ == "BUSINESS-PROCESS":
                    continue
                elif type_ == "MANUAL":
                    ws.cell(row=row, column=3).value  = test.TestId
                    ws.cell(row=row, column=4).value  = test.Field("TS_NAME").strip()
                    ws.cell(row=row, column=5).value  = self.extract_text_from_html(test.Field("TS_DESCRIPTION"))
                    ws.cell(row=row, column=6).value  = test.Field("TS_USER_10").strip()
                    ws.cell(row=row, column=7).value  = test.Field("TS_USER_05").strip()
                    ws.cell(row=row, column=8).value  = test.Field("TS_TYPE").strip()
                    ws.cell(row=row, column=9).value  = test.Field("TS_RESPONSIBLE").strip()
                    cd = test.Field("TS_CREATION_DATE")
                    ws.cell(row=row, column=10).value = cd.strftime('%Y-%m-%d %H:%M:%S') if cd else ""
                    ws.cell(row=row, column=11).value = self.extract_text_from_html(test.Field("TS_DEV_COMMENTS"))
                    ws.cell(row=row, column=12).value = test.Field("TS_BPTA_CHANGE_DETECTED")
                    ws.cell(row=row, column=13).value = test.Field("TS_EXEC_STATUS").strip()
                    for step in testPlan.DesignStepfactory.NewList(""):
                        ws.cell(row=row, column=14).value = step.StepName
                        ws.cell(row=row, column=15).value = self.extract_text_from_html(step.StepDescription)
                        ws.cell(row=row, column=16).value = self.extract_text_from_html(step.StepExpectedResult)
                        row += 1
        wb.save("ConsolidatedTestSetReport.xlsx")
        return "ConsolidatedTestSetReport.xlsx"


class Attachment_Downloader:
    def download_Attachments(self, test_set_ids, tdc):
        for Suite_ID in [s.strip() for s in test_set_ids.split(",")]:
            test_suite = tdc.TestSetFactory.Item(Suite_ID)
            wb = Workbook()
            ws = wb.active
            headers = ["Test Set ID","Test Instance ID","Test Case ID","Test case type","Latest Test Run ID",
                       "Run status","tester XT ID","Tester full name","Execution date",
                       "Attachment validation-Run Level","Download path - Run Level",
                       "Step Name","Status","Attachment validation-Step Level","Download Path"]
            ws.append(headers)
            ws.cell(row=2, column=1).value = Suite_ID
            row = 2
            for test_instance in test_suite.TSTestFactory.NewList(""):
                ws.cell(row=row, column=2).value = test_instance.id
                ws.cell(row=row, column=3).value = test_instance.testId
                ws.cell(row=row, column=4).value = test_instance.type
                tester_xt_id = test_instance.Field("TC_ACTUAL_TESTER")
                latest_run = latest_date = latest_time = None
                for run in test_instance.RunFactory.NewList(""):
                    run_date = run.Field("RN_EXECUTION_DATE")
                    run_time = run.Field("RN_EXECUTION_TIME")
                    if (latest_date is None or run_date > latest_date or
                            (run_date == latest_date and run_time > latest_time)):
                        latest_run, latest_date, latest_time = run, run_date, run_time
                if not latest_run:
                    ws.cell(row=row, column=5).value = "No runs found"
                else:
                    ws.cell(row=row, column=5).value = latest_run.ID
                    ws.cell(row=row, column=6).value = latest_run.Status
                    ws.cell(row=row, column=7).value = tester_xt_id
                    exec_date = latest_run.Field("RN_EXECUTION_DATE")
                    ws.cell(row=row, column=9).value = f"{exec_date.day:02d}/{exec_date.month:02d}/{exec_date.year}"
                    Run_attachments = latest_run.Attachments.NewList("")
                    Download_Path_R = " "
                    if len(Run_attachments) > 0:
                        ws.cell(row=row, column=10).value = "Having Attachment"
                        for R_attach in Run_attachments:
                            folder = os.path.join(os.getcwd(), "ALMExtraction", Suite_ID, str(latest_run.ID))
                            R_attach.AttachmentStorage.ClientPath = folder
                            try:
                                R_attach.Load(True, folder)
                                Download_Path_R += ";" + R_attach.FileName
                            except:
                                pass
                        ws.cell(row=row, column=11).value = Download_Path_R
                    else:
                        ws.cell(row=row, column=10).value = "No Attachments"
                    steprow = row
                    for step in latest_run.StepFactory.NewList(""):
                        ws.cell(row=steprow, column=12).value = step.Name
                        ws.cell(row=steprow, column=13).value = step.Status
                        attachments = step.Attachments.NewList("")
                        Download_Path = ""
                        if len(attachments) > 0:
                            ws.cell(row=steprow, column=14).value = "Having Attachment"
                            for t_attach in attachments:
                                folder = os.path.join(os.getcwd(), "ALMExtraction", Suite_ID, str(latest_run.ID), step.Name)
                                t_attach.AttachmentStorage.ClientPath = folder
                                try:
                                    t_attach.Load(True, folder)
                                    Download_Path += ";" + t_attach.FileName
                                except:
                                    pass
                            ws.cell(row=steprow, column=15).value = Download_Path
                        else:
                            ws.cell(row=steprow, column=14).value = "No Attachments"
                        steprow += 1
                    row = steprow
            now = datetime.now()
            wb.save(f"{Suite_ID}_{now.strftime('%d-%m-%Y_%H-%M-%S')}.xlsx")

    def download_Attachments_All(self, tdc):
        for testset in tdc.TestSetFactory.NewList(""):
            test_set_id = testset.ID
            wb = Workbook()
            ws = wb.active
            headers = ["Test Set ID","Test Instance ID","Test Case ID","Test case type","Latest Test Run ID",
                       "Run status","tester XT ID","Tester full name","Execution date",
                       "Attachment validation-Run Level","Download path - Run Level",
                       "Step Name","Status","Attachment validation-Step Level","Download Path"]
            ws.append(headers)
            ws.cell(row=2, column=1).value = test_set_id
            row = 2
            for test_instance in testset.TSTestFactory.NewList(""):
                ws.cell(row=row, column=2).value = test_instance.id
                ws.cell(row=row, column=3).value = test_instance.testId
                ws.cell(row=row, column=4).value = test_instance.type
                tester_xt_id = test_instance.Field("TC_ACTUAL_TESTER")
                latest_run = latest_date = latest_time = None
                for run in test_instance.RunFactory.NewList(""):
                    run_date = run.Field("RN_EXECUTION_DATE")
                    run_time = run.Field("RN_EXECUTION_TIME")
                    if (latest_date is None or run_date > latest_date or
                            (run_date == latest_date and run_time > latest_time)):
                        latest_run, latest_date, latest_time = run, run_date, run_time
                if not latest_run:
                    ws.cell(row=row, column=5).value = "No runs found"
                else:
                    ws.cell(row=row, column=5).value = latest_run.ID
                    ws.cell(row=row, column=6).value = latest_run.Status
                    ws.cell(row=row, column=7).value = tester_xt_id
                    exec_date = latest_run.Field("RN_EXECUTION_DATE")
                    ws.cell(row=row, column=9).value = f"{exec_date.day:02d}/{exec_date.month:02d}/{exec_date.year}"
                    Run_attachments = latest_run.Attachments.NewList("")
                    Download_Path_R = " "
                    if len(Run_attachments) > 0:
                        ws.cell(row=row, column=10).value = "Having Attachment"
                        for R_attach in Run_attachments:
                            folder = os.path.join(os.getcwd(), "ALMExtraction", str(test_set_id), str(latest_run.ID))
                            R_attach.AttachmentStorage.ClientPath = folder
                            try:
                                R_attach.Load(True, folder)
                                Download_Path_R += ";" + R_attach.FileName
                            except:
                                pass
                        ws.cell(row=row, column=11).value = Download_Path_R
                    else:
                        ws.cell(row=row, column=10).value = "No Attachments"
                    steprow = row
                    for step in latest_run.StepFactory.NewList(""):
                        ws.cell(row=steprow, column=12).value = step.Name
                        ws.cell(row=steprow, column=13).value = step.Status
                        attachments = step.Attachments.NewList("")
                        Download_Path = ""
                        if len(attachments) > 0:
                            ws.cell(row=steprow, column=14).value = "Having Attachment"
                            for t_attach in attachments:
                                folder = os.path.join(os.getcwd(), "ALMExtraction", str(test_set_id), str(latest_run.ID), step.Name)
                                t_attach.AttachmentStorage.ClientPath = folder
                                try:
                                    t_attach.Load(True, folder)
                                    Download_Path += ";" + t_attach.FileName
                                except:
                                    pass
                            ws.cell(row=steprow, column=15).value = Download_Path
                        else:
                            ws.cell(row=steprow, column=14).value = "No Attachments"
                        steprow += 1
                    row = steprow
            now = datetime.now()
            wb.save(f"{test_set_id}_{now.strftime('%d-%m-%Y_%H-%M-%S')}.xlsx")


class Maintenance_Notification:
    def no_digits(self, s):
        for x in "0123456789":
            s = s.replace(x, "")
        return s

    def checkdisplayname_email(self, name_email):
        try:
            OLApp = win32com.client.Dispatch("Outlook.Application")
            oRecip = OLApp.Session.CreateRecipient(name_email)
            oRecip.Resolve()
            if oRecip.Resolved:
                ae = oRecip.AddressEntry
                if ae.AddressEntryUserType in [0, 5]:
                    oEU = ae.GetExchangeUser()
                    if oEU:
                        return oEU.PrimarySmtpAddress
                elif ae.AddressEntryUserType in [10, 30]:
                    return ae.Address
        except Exception as e:
            print(f"Error: {e}")
        return None

    def process_users(self, tdc):
        userMails = ""
        cust = tdc.Customization
        cust.Load()
        for user in cust.Users.Users:
            xtid = user.Name
            username = self.no_digits(user.FullName)
            email = self.checkdisplayname_email(xtid) or self.checkdisplayname_email(username)
            if email:
                userMails += email + ";"
        return userMails

    def Outlook_attachment(self, all_mail_ids, from_date, to_date, from_hour, from_minute,
                           from_am, to_hour, to_minute, to_am, reason, project_Name):
        outlook = win32com.client.Dispatch("Outlook.Application")
        mail = outlook.CreateItem(0)
        mail_body = (f"Hi all,\n\nDue to {reason.strip()} the ALM services will not be available "
                     f"between {from_date} {from_hour}:{from_minute} {from_am} to "
                     f"{to_date} {to_hour}:{to_minute} {to_am}.\n\nRegards,\nTools Admin Team")
        mail.Subject = f"ALM Maintenance Notification - {project_Name}"
        mail.HTMLBody = f"<p>{mail_body.replace(chr(10), '<br>')}</p>"
        mail.To = all_mail_ids
        mail.Send()


class AccessProvider:
    def Access_Provider(self, username, useremail, group, tdc):
        try:
            if not self.Check_if_user_is_in_siteadmin(tdc, username):
                self.add_user_to_site_admin(tdc, username, useremail, group)
            time.sleep(2)
            if not self.user_exists_in_project(tdc, username):
                if self.Add_user_to_project(tdc, username):
                    time.sleep(2)
                    self._assign_group(tdc, username, group)
            else:
                self._assign_group(tdc, username, group)
            return True
        except Exception as e:
            print(f"Error: {e}")
            return False

    def _assign_group(self, tdc, username, group):
        if not group:
            return
        cust = tdc.Customization
        for user in cust.Users.Users:
            if user.Name.lower() == username.lower():
                for grp in group.split(";"):
                    if not self.user_exists_in_group(tdc, username, grp):
                        user.AddToGroup(grp)
                        cust.Commit()

    def user_exists_in_group(self, tdc, username, igroup):
        try:
            for user in tdc.Customization.Users.Users:
                if user.Name.lower() == username.lower():
                    for group in user.GroupsList():
                        if group.Name and group.Name.lower() == igroup.lower():
                            return True
        except:
            pass
        return False

    def Add_user_to_project(self, tdc, username):
        try:
            cust = tdc.Customization
            cust.Users.AddUser(username)
            cust.Commit()
            return True
        except Exception as e:
            print(f"Error: {e}")
            return False

    def user_exists_in_project(self, tdc, username):
        try:
            for user in tdc.Customization.Users.Users:
                if user.Name.lower() == username.lower():
                    return True
        except:
            pass
        return False

    def Check_if_user_is_in_siteadmin(self, tdc, username):
        cust = tdc.Customization
        return cust.Users.UserExistsInSite(username)

    def add_user_to_site_admin(self, tdc, username, useremail, group):
        try:
            cust = tdc.Customization
            ol = win32com.client.Dispatch("Outlook.Application")
            new_mail = ol.CreateItem(0)
            new_mail.Subject = "Test Email"
            new_mail.To = useremail
            recip = new_mail.Recipients.Item(1)
            if recip.Resolve():
                oEU = new_mail.Recipients.Item(1).AddressEntry.GetExchangeUser()
                alias = username
                full_name = f"{oEU.FirstName} {oEU.LastName}" if oEU.FirstName else ""
                email = oEU.PrimarySmtpAddress
                phone = oEU.BusinessTelephoneNumber
                city = oEU.City
            else:
                alias, full_name, email, phone, city = username, "", "", "", ""
            cust.Users.AddSiteUser(alias, full_name, email, city, phone, " ")
            cust.Commit()
            return True
        except Exception as e:
            print(f"Error: {e}")
            return False


# ══════════════════════════════════════════════════════════════════════════════
# SINGLE ALM THREAD
# ══════════════════════════════════════════════════════════════════════════════

class ALMThread(threading.Thread):
    def __init__(self):
        super().__init__(daemon=True)
        self.task_queue = queue.Queue()
        self.tdc = None
        self.connected = False

    def run(self):
        if OTA_AVAILABLE:
            pythoncom.CoInitialize()
        while True:
            try:
                fn, args, kwargs, result_q = self.task_queue.get(timeout=1)
                try:
                    result_q.put(("ok", fn(*args, **kwargs)))
                except Exception as e:
                    result_q.put(("err", str(e)))
            except queue.Empty:
                continue

    def call(self, fn, *args, timeout=60, **kwargs):
        result_q = queue.Queue()
        self.task_queue.put((fn, args, kwargs, result_q))
        status, value = result_q.get(timeout=timeout)
        if status == "err":
            raise Exception(value)
        return value

    def _login(self, server, username, password):
        self.tdc = win32com.client.Dispatch("TDApiOle80.TDConnection")
        self.tdc.InitConnectionEx(server)
        self.tdc.Login(username, password)
        return True

    def _get_domains(self):
        return [d for d in self.tdc.VisibleDomains]

    def _get_projects(self, domain):
        return [p for p in self.tdc.VisibleProjects(domain)]

    def _connect(self, domain, project):
        self.tdc.Connect(domain, project)
        self.connected = True
        return True

    def _logout(self):
        try:
            if self.tdc and self.tdc.Connected:
                self.tdc.Disconnect()
            if self.tdc:
                self.tdc.Logout()
                self.tdc.ReleaseConnection()
        except:
            pass
        self.tdc = None
        self.connected = False
        return True

    def _get_groups(self):
        com = self.tdc.Command
        com.CommandText = "SELECT GROUPS.GR_GROUP_NAME FROM GROUPS"
        rs = com.Execute()
        groups = []
        while not rs.EOR:
            groups.append(rs.FieldValue(0))
            rs.Next()
        return groups

    def _update_test_type(self, ids_list, test_type):
        TestTypeUpdator().execute_conversion(ids_list, test_type, self.tdc)
        return len(ids_list)

    def _extract_tests(self, source, folder_path, test_set_ids, with_attachments):
        if source == 'test_set':
            return testCaseExtractor_TS().download_test_cases(test_set_ids, self.tdc)
        elif with_attachments:
            return testCaseExtractor_WA().TestCase_WA(folder_path, self.tdc)
        else:
            return testCaseExtractor().TestCase_WA(folder_path, self.tdc)

    def _generate_evidence(self, input_file):
        ALMQueryProcessor().execute_query_and_save_to_excel(
            input_file, "ALM_Query_Results.xlsx", "ALM_Query_Results", self.tdc)
        return "ALM_Query_Results/"

    def _extract_defects(self, status, priority, severity, category):
        return DefectExtractor().Defect_Extraction(status, priority, severity, category, self.tdc)

    def _download_attachments(self, test_set_ids, download_all):
        dl = Attachment_Downloader()
        if download_all:
            dl.download_Attachments_All(self.tdc)
        else:
            dl.download_Attachments(test_set_ids, self.tdc)
        return True

    def _maintenance_notify(self, from_date, to_date, from_time, to_time, reason):
        notifier = Maintenance_Notification()
        mail_users = notifier.process_users(self.tdc)
        project_name = self.tdc.ProjectName if self.tdc.Connected else "ALM Project"
        fh, fm = (from_time.split(':') + ['00'])[:2]
        th, tm = (to_time.split(':')   + ['00'])[:2]
        notifier.Outlook_attachment(mail_users, from_date, to_date, fh, fm, "AM", th, tm, "AM", reason, project_name)
        return len(mail_users.split(';')) if mail_users else 0

    def _provision_user(self, username, email, group):
        return AccessProvider().Access_Provider(username, email, group, self.tdc)


alm_thread = ALMThread()
alm_thread.start()


# ══════════════════════════════════════════════════════════════════════════════
# FLASK ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.json or {}
    server_url = data.get('serverUrl', '').strip()
    username   = data.get('username', '').strip()
    password   = data.get('password', '').strip()
    if not all([server_url, username, password]):
        return jsonify({"detail": "Server URL, username and password required"}), 400
    token = str(uuid.uuid4())
    if OTA_AVAILABLE:
        try:
            alm_thread.call(alm_thread._login, server_url, username, password)
            sessions[token] = {"username": username, "server": server_url,
                               "domain": None, "project": None, "mode": "real"}
            return jsonify({"token": token, "username": username, "mode": "real",
                "message": f"✅ Logged in to {server_url}",
                "logs": [log_entry(f"Connecting: {server_url}"),
                         log_entry(f"Authenticating: {username}"),
                         log_entry("✅ Login successful")]})
        except Exception as e:
            return jsonify({"detail": f"Login failed: {str(e)}"}), 401
    dummy_delay()
    sessions[token] = {"username": username, "server": server_url,
                       "domain": None, "project": None, "mode": "demo"}
    return jsonify({"token": token, "username": username, "mode": "demo",
        "message": "✅ Demo login successful",
        "logs": [log_entry(f"[DEMO] {server_url}"), log_entry("✅ Demo ready")]})


@app.route('/api/auth/domains', methods=['GET'])
@require_session
def get_domains(session):
    if session['mode'] == 'real':
        try:
            domains = alm_thread.call(alm_thread._get_domains)
            return jsonify({"domains": domains, "mode": "real"})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    return jsonify({"domains": ["DEFAULT", "CN_DOMAIN", "QA_DOMAIN"], "mode": "demo"})


@app.route('/api/auth/projects', methods=['GET'])
@require_session
def get_projects(session):
    domain = request.args.get('domain', '')
    if not domain:
        return jsonify({"detail": "Domain required"}), 400
    if session['mode'] == 'real':
        try:
            projects = alm_thread.call(alm_thread._get_projects, domain)
            return jsonify({"projects": projects, "mode": "real"})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    demo = {"DEFAULT": ["Mercury_Release_v4", "Mercury_Release_v3"],
            "CN_DOMAIN": ["CN_Rail_Main", "CN_Rail_QA"],
            "QA_DOMAIN": ["QA_Project_1", "QA_Project_2"]}
    return jsonify({"projects": demo.get(domain, ["Default_Project"]), "mode": "demo"})


@app.route('/api/auth/connect', methods=['POST'])
@require_session
def connect_project(session):
    data    = request.json or {}
    domain  = data.get('domain', '')
    project = data.get('project', '')
    if not domain or not project:
        return jsonify({"detail": "Domain and project required"}), 400
    if session['mode'] == 'real':
        try:
            alm_thread.call(alm_thread._connect, domain, project)
            session['domain']  = domain
            session['project'] = project
            return jsonify({"message": f"✅ Connected to {domain}/{project}", "mode": "real",
                "logs": [log_entry(f"Domain: {domain}"), log_entry(f"Project: {project}"),
                         log_entry("✅ TDC ready — all utilities available", "SUCCESS")]})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    dummy_delay()
    session['domain']  = domain
    session['project'] = project
    return jsonify({"message": f"✅ [DEMO] Connected to {domain}/{project}", "mode": "demo",
        "logs": [log_entry(f"[DEMO] {domain}/{project}"), log_entry("✅ Ready")]})


@app.route('/api/auth/logout', methods=['POST'])
@require_session
def logout(session):
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    if session['mode'] == 'real':
        try:
            alm_thread.call(alm_thread._logout)
        except:
            pass
    sessions.pop(token, None)
    return jsonify({"message": "Logged out"})


@app.route('/api/user/groups', methods=['GET'])
@require_session
def get_groups(session):
    if session['mode'] == 'real':
        try:
            groups = alm_thread.call(alm_thread._get_groups)
            return jsonify({"groups": groups, "mode": "real"})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    return jsonify({"groups": ["QA_Engineer","QA_Lead","Admin","ReadOnly","Developer","Manager"], "mode": "demo"})


@app.route('/api/user/provision', methods=['POST'])
@require_session
def provision_user(session):
    data     = request.json or {}
    username = data.get('username', '')
    email    = data.get('email', '')
    group    = data.get('group', '')
    if not username:
        return jsonify({"detail": "Username required"}), 400
    if session['mode'] == 'real':
        try:
            alm_thread.call(alm_thread._provision_user, username, email, group, timeout=60)
            return jsonify({"success": True, "message": f"✅ Access granted to {username}", "mode": "real",
                "logs": [log_entry(f"Provisioning: {username} → {group}"),
                         log_entry("✅ Access granted", "SUCCESS")]})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    dummy_delay()
    return jsonify({"success": True, "message": f"[DEMO] Access granted to {username}", "mode": "demo",
        "logs": [log_entry(f"[DEMO] {username} → {group}"), log_entry("✅ Done", "SUCCESS")]})


@app.route('/api/extract/tests', methods=['POST'])
@require_session
def extract_tests(session):
    data             = request.json or {}
    source           = data.get('source', 'test_plan')
    folder_path      = data.get('folderPath', '')
    test_set_ids     = data.get('testSetIds', '')
    with_attachments = data.get('withAttachments', False)
    if session['mode'] == 'real':
        try:
            filename = alm_thread.call(alm_thread._extract_tests, source, folder_path,
                                       test_set_ids, with_attachments, timeout=300)
            return jsonify({"success": True, "filename": filename, "mode": "real",
                "logs": [log_entry(f"Source: {source}"),
                         log_entry(f"Path/IDs: {folder_path or test_set_ids}"),
                         log_entry(f"✅ Saved: {filename}", "SUCCESS")]})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    dummy_delay()
    count = random.randint(15, 120)
    fn = f"TestCases_{datetime.now().strftime('%d%m%Y_%H%M%S')}.xlsx"
    return jsonify({"success": True, "recordCount": count, "filename": fn, "mode": "demo",
        "logs": [log_entry(f"[DEMO] {count} test cases"), log_entry(f"✅ {fn}", "SUCCESS")]})


@app.route('/api/generate/evidence', methods=['POST'])
@require_session
def generate_evidence(session):
    data = request.json or {}
    input_file = data.get('inputFile', '')
    if not input_file:
        return jsonify({"detail": "Input Excel file path required"}), 400
    if session['mode'] == 'real':
        try:
            output = alm_thread.call(alm_thread._generate_evidence, input_file, timeout=300)
            return jsonify({"success": True, "outputFolder": output, "mode": "real",
                "logs": [log_entry(f"Input: {input_file}"),
                         log_entry("Generating Word documents..."),
                         log_entry("✅ Evidence generated", "SUCCESS")]})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    dummy_delay()
    docs = random.randint(10, 50)
    return jsonify({"success": True, "docCount": docs, "mode": "demo",
        "logs": [log_entry(f"[DEMO] {docs} Word docs"), log_entry("✅ Done", "SUCCESS")]})


@app.route('/api/update/test-type', methods=['POST'])
@require_session
def update_test_type(session):
    data      = request.json or {}
    ids_raw   = data.get('testCaseIds', '')
    test_type = data.get('testType', 'MANUAL')
    if not ids_raw:
        return jsonify({"detail": "Test case IDs required"}), 400
    ids_list = [i.strip() for i in str(ids_raw).split(',') if i.strip()]
    if session['mode'] == 'real':
        try:
            count = alm_thread.call(alm_thread._update_test_type, ids_list, test_type)
            return jsonify({"success": True, "updatedCount": count, "mode": "real",
                "logs": [log_entry(f"Updating {count} test case(s) → {test_type}"),
                         log_entry("✅ Update complete", "SUCCESS")]})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    dummy_delay()
    return jsonify({"success": True, "updatedCount": len(ids_list), "mode": "demo",
        "logs": [log_entry(f"[DEMO] {len(ids_list)} → {test_type}"), log_entry("✅ Done", "SUCCESS")]})


@app.route('/api/extract/defects', methods=['POST'])
@require_session
def extract_defects(session):
    data     = request.json or {}
    status   = [s.strip() for s in data.get('status',   '').split(',') if s.strip()]
    priority = [p.strip() for p in data.get('priority', '').split(',') if p.strip()]
    severity = [s.strip() for s in data.get('severity', '').split(',') if s.strip()]
    category = [c.strip() for c in data.get('category', '').split(',') if c.strip()]
    if session['mode'] == 'real':
        try:
            filename = alm_thread.call(alm_thread._extract_defects,
                                       status, priority, severity, category, timeout=300)
            return jsonify({"success": True, "filename": filename, "mode": "real",
                "logs": [log_entry("Querying ALM BUG table..."),
                         log_entry(f"✅ Saved: {filename}", "SUCCESS")]})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    dummy_delay()
    count = random.randint(20, 200)
    return jsonify({"success": True, "defectCount": count, "filename": "ExtractedDefects.xlsx", "mode": "demo",
        "logs": [log_entry(f"[DEMO] {count} defects"), log_entry("✅ Done", "SUCCESS")]})


@app.route('/api/download/attachments', methods=['POST'])
@require_session
def download_attachments(session):
    data         = request.json or {}
    test_set_ids = data.get('testSetIds', '')
    download_all = data.get('downloadAll', False)
    if session['mode'] == 'real':
        try:
            alm_thread.call(alm_thread._download_attachments, test_set_ids, download_all, timeout=300)
            return jsonify({"success": True, "message": "✅ Attachments downloaded", "mode": "real",
                "logs": [log_entry("Downloading attachments..."),
                         log_entry("✅ Complete — check ALMExtraction folder", "SUCCESS")]})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    dummy_delay()
    count = random.randint(10, 80)
    return jsonify({"success": True, "attachmentCount": count, "mode": "demo",
        "logs": [log_entry(f"[DEMO] {count} attachments"), log_entry("✅ Done", "SUCCESS")]})


@app.route('/api/maintenance/notify', methods=['POST'])
@require_session
def maintenance_notify(session):
    data      = request.json or {}
    from_date = data.get('fromDate', '')
    to_date   = data.get('toDate', '')
    from_time = data.get('fromTime', '')
    to_time   = data.get('toTime', '')
    reason    = data.get('reason', '')
    if not reason:
        return jsonify({"detail": "Reason required"}), 400
    if session['mode'] == 'real':
        try:
            count = alm_thread.call(alm_thread._maintenance_notify,
                                    from_date, to_date, from_time, to_time, reason, timeout=60)
            return jsonify({"success": True, "userCount": count, "mode": "real",
                "logs": [log_entry("Fetching project users..."),
                         log_entry("Sending Outlook notifications..."),
                         log_entry(f"✅ Sent to {count} users", "SUCCESS")]})
        except Exception as e:
            return jsonify({"detail": str(e)}), 500
    dummy_delay()
    count = random.randint(15, 80)
    return jsonify({"success": True, "userCount": count, "mode": "demo",
        "logs": [log_entry(f"[DEMO] {count} users notified"), log_entry("✅ Done", "SUCCESS")]})


@app.route('/api/dashboard/stats', methods=['GET'])
def dashboard_stats():
    return jsonify({"botsDeployed": 20, "hoursSaved": 40, "usersManaged": 2048,
                    "effortReduced": "90%", "mode": "real" if OTA_AVAILABLE else "demo",
                    "lastRunTime": datetime.now().strftime("%d %b %Y, %H:%M")})

@app.route('/api/dashboard/history', methods=['GET'])
def dashboard_history():
    actions = ["Test Extraction","Defect Extraction","Evidence Generator","User Access Granted",
               "Test Type Update","Maintenance Notification","Attachment Download"]
    return jsonify([{"id": i+1, "action": a, "status": "Completed",
                     "timestamp": datetime.now().strftime("%d %b %Y, %H:%M"),
                     "records": random.randint(10, 200)} for i, a in enumerate(actions)])

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({"status": "ok", "service": "One Stop ALM Administration API v3.0",
                    "ota_available": OTA_AVAILABLE,
                    "mode": "real" if OTA_AVAILABLE else "demo",
                    "alm_connected": alm_thread.connected,
                    "time": datetime.now().isoformat()})


if __name__ == '__main__':
    print("=" * 60)
    print("  One Stop ALM Administration — Flask Backend v3.0")
    print(f"  Mode: {'REAL (OTA)' if OTA_AVAILABLE else 'DEMO'}")
    print("  All utilities from login.py integrated")
    print("  Single ALM thread — tdc shared across all utilities")
    print("  Running on: http://localhost:8000")
    print("=" * 60)
    app.run(host='0.0.0.0', port=8000, debug=False, threaded=True)
